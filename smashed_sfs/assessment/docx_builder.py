import io
import re

import docx
from docx.shared import Pt

# Matches **word** markup the AI uses to mark a negation (e.g. **NOT**) so
# it can be rendered bold - see assessment/prompts.py's instruction to the
# model. Kept as its own constant since both the question-stem renderer
# below and (if ever needed) other renderers should agree on the markup.
BOLD_MARKUP = re.compile(r'\*\*(.+?)\*\*')


def _add_text_with_bold_markup(paragraph, text, *, bold_all=False):
    """Splits `text` on **bold** markup and adds each piece as its own
    run, bolding the marked pieces (or every piece, when the whole line -
    e.g. a question stem - should be bold regardless of markup)."""
    pos = 0
    for match in BOLD_MARKUP.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()]).bold = bold_all
        run = paragraph.add_run(match.group(1))
        run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:]).bold = bold_all


def build_assessment_docx(*, assessment, teacher_name, learning_area, grade_section):
    """assessment is the AI's parsed JSON: assessment_title plus up to
    three item arrays (multiple_choice, true_false_plain,
    true_false_modified - any of which may be empty/absent). Builds the
    flowing document format from the tool's own prompt spec (title,
    instructions, continuously-numbered questions across parts, then a
    single Answer Key & Taxonomy section) rather than a branded slide
    deck - see the tool's design notes for why that's out of scope here.
    Returns raw docx bytes; nothing is written to disk."""
    mc_items = assessment.get('multiple_choice') or []
    tf_plain_items = assessment.get('true_false_plain') or []
    tf_modified_items = assessment.get('true_false_modified') or []

    document = docx.Document()

    document.add_heading(assessment.get('assessment_title', 'Assessment'), level=1)
    meta = document.add_paragraph()
    meta.add_run(f'Learning Area: {learning_area}    Grade & Section: {grade_section}\n').italic = True
    meta.add_run(f'Prepared by: {teacher_name}').italic = True

    document.add_paragraph(
        'Instructions: Read each item carefully. For Multiple Choice, choose the letter of the best answer. '
        'For True or False, write TRUE or FALSE; where a correction is asked for, explain what would make a '
        'false statement true.'
    )

    item_number = 1
    answer_key_rows = []  # (item_number, part_label, answer_text, bloom_level)

    if mc_items:
        document.add_heading(f'Part I: Multiple Choice ({len(mc_items)} items)', level=2)
        for item in mc_items:
            p = document.add_paragraph()
            p.add_run(f'{item_number}. ')
            _add_text_with_bold_markup(p, item.get('question', ''))
            for letter in ('a', 'b', 'c', 'd'):
                document.add_paragraph(f"   {letter}. {item.get('choices', {}).get(letter, '')}")
            answer_key_rows.append((item_number, 'MC', item.get('correct', ''), item.get('bloom_level', '')))
            item_number += 1

    if tf_plain_items:
        document.add_heading(f'Part II: True or False ({len(tf_plain_items)} items)', level=2)
        for item in tf_plain_items:
            document.add_paragraph(f"{item_number}. {item.get('statement', '')}")
            answer_key_rows.append((item_number, 'TF', item.get('correct', ''), item.get('bloom_level', '')))
            item_number += 1

    if tf_modified_items:
        document.add_heading(f'Part III: Modified True or False ({len(tf_modified_items)} items)', level=2)
        document.add_paragraph(
            'If your answer is FALSE, explain what should be corrected to make the statement true.'
        )
        for item in tf_modified_items:
            document.add_paragraph(f"{item_number}. {item.get('statement', '')}")
            answer_key_rows.append((
                item_number, 'Modified TF', item.get('correct', ''), item.get('bloom_level', ''),
                item.get('correction', '') if item.get('correct') == 'FALSE' else '',
            ))
            item_number += 1

    document.add_paragraph()
    document.add_heading('--- Answer Key & Taxonomy ---', level=2)
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text, header[3].text = (
        'Item No.', 'Part', 'Correct Answer / Correction', "Bloom's Level"
    )
    for cell in header:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    for row_data in answer_key_rows:
        number, part, correct, bloom_level, *rest = row_data
        correction = rest[0] if rest else ''
        answer_text = f'{correct} - {correction}' if correction else correct
        row = table.add_row().cells
        row[0].text = str(number)
        row[1].text = part
        row[2].text = answer_text
        row[3].text = bloom_level
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
