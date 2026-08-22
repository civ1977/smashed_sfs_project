import io

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# DepEd Order No. 003, s. 2026, Annex A - "FOR TEACHER AI USE DECLARATION."
# Filled in automatically from the teacher's account name and this
# generation's own inputs, rather than left for the teacher to write by
# hand - see grades/views.py-style constant comments for why this is kept
# as one clearly-labeled source rather than inlined.
AI_DECLARATION_TEMPLATE = (
    'Consistent with the policy guidelines on the use of AI in basic education, I, {teacher_name}, '
    'hereby declare that I have used AI tools to assist in the preparation and delivery of teaching '
    'and learning materials.\n\n'
    'I used {provider} to draft the lesson plan titled "{lesson_title}" for the purposes of lesson '
    'planning. The prompts used were: competencies - {competencies}; additional instructions - '
    '{added_instructions}.\n\n'
    'I affirm that the use of AI tools was intended solely to enhance the quality of instructional '
    'material, align with curriculum standards, and support the teaching and learning process. All '
    'outputs have been reviewed, adapted, edited, and validated using my professional expertise and '
    'agency to ensure accuracy, appropriateness, and alignment with learners’ developmental '
    'needs/levels.\n\n'
    'Likewise, I affirm that no confidential learner information or sensitive institutional '
    'information data was shared with the AI platform/app identified above in the process.'
)

FIELD_ROWS = [
    ('learning_competency', 'Learning Competency and Curriculum Standards'),
    ('learning_objectives', 'Learning Objectives'),
    ('learner_context', 'Learner Context'),
    ('pre_lesson', 'Pre-Lesson'),
    ('flow', 'Flow'),
    ('learning_resources', 'Learning Resources'),
    ('integration_and_contextualization', 'Opportunities for Integration and Contextualization'),
    ('formative_assessment', 'Formative Assessment'),
    ('extended_learning', 'Extended Learning Opportunities'),
]


def build_ai_declaration(*, teacher_name, provider, lesson_title, competencies, added_instructions):
    return AI_DECLARATION_TEMPLATE.format(
        teacher_name=teacher_name,
        provider=provider,
        lesson_title=lesson_title,
        competencies=competencies.strip() or 'N/A',
        added_instructions=added_instructions.strip() or 'N/A',
    )


def _set_cell_text(cell, text, bold=False):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def build_lesson_plan_docx(*, plan, header_fields, ai_declaration):
    """plan is the AI's parsed JSON (lesson_title + days[]); header_fields
    is a dict of the info-table values (learning_area, teacher_name,
    grade_section, teaching_dates, num_sessions, references). No school
    letterhead/page header-footer and no signature block are included -
    see the Lesson Planning tool's own page copy for why (an AI can't
    know who a school's Master Teacher or Principal is, and sign-off is a
    physical step that happens after this document is downloaded).
    Returns raw docx bytes; nothing is written to disk."""
    document = docx.Document()

    document.add_heading(plan.get('lesson_title', 'Lesson Plan'), level=1)

    info_table = document.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    info_rows = [
        ('Lesson Title', plan.get('lesson_title', '')),
        ('Learning Area', header_fields['learning_area']),
        ('Name of Teacher', header_fields['teacher_name']),
        ('Grade Level and Section', header_fields['grade_section']),
        ('Teaching Dates', header_fields['teaching_dates']),
        ('No. of Sessions', str(header_fields['num_sessions'])),
        ('References', header_fields.get('references') or 'N/A'),
        ('Declaration of AI Use', ai_declaration),
    ]
    for label, value in info_rows:
        row = info_table.add_row()
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[1], value)

    document.add_paragraph()

    days = plan.get('days', [])
    main_table = document.add_table(rows=0, cols=1 + len(days))
    main_table.style = 'Table Grid'

    header_row = main_table.add_row()
    _set_cell_text(header_row.cells[0], '')
    for i, day in enumerate(days):
        _set_cell_text(header_row.cells[i + 1], day.get('day_label', f'Day {i + 1}'), bold=True)
        header_row.cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for field_key, field_label in FIELD_ROWS:
        row = main_table.add_row()
        _set_cell_text(row.cells[0], field_label, bold=True)
        for i, day in enumerate(days):
            value = day.get(field_key, '')
            if isinstance(value, list):
                value = '\n'.join(f'{idx}. {item}' for idx, item in enumerate(value, start=1))
            _set_cell_text(row.cells[i + 1], value)

    reflections_row = main_table.add_row()
    _set_cell_text(reflections_row.cells[0], 'Reflections', bold=True)
    for i in range(len(days)):
        _set_cell_text(reflections_row.cells[i + 1], '')

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
